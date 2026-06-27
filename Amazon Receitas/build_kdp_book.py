import argparse
import os
import re
import sys
from io import BytesIO
from pathlib import Path

import fitz
import requests
from PIL import Image
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches


def extract_pdf_text(pdf_path: Path) -> str:
    doc = fitz.open(pdf_path)
    pages = []
    for page in doc:
        text = page.get_text("text")
        pages.append(text)
    return "\n\n".join(pages)


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def parse_page_header_category(line: str) -> str | None:
    if not line:
        return None
    normalized = line.replace("–", "-").replace("—", "-").replace("•", "-").strip()
    if "200 RECEITAS LOWCARB" not in normalized.upper():
        return None
    match = re.search(r"200\s+RECEITAS\s+LOWCARB\s*[-]\s*(.+)$", normalized, re.I)
    if not match:
        return None
    return match.group(1).strip().title()


def is_pdf_header_line(line: str) -> bool:
    if not line:
        return False
    upper = line.upper()
    if "200 RECEITAS LOWCARB" in upper:
        return True
    if line.isdigit():
        return True
    return False


def split_into_recipes(text: str) -> list[dict]:
    raw_lines = [line.strip() for line in normalize_text(text).splitlines()]
    parsed_lines: list[tuple[str, str | None]] = []
    current_category: str | None = None

    for line in raw_lines:
        if not line:
            continue
        if is_pdf_header_line(line):
            category = parse_page_header_category(line)
            if category:
                current_category = category
            continue
        if line.isdigit():
            continue
        parsed_lines.append((line, current_category))

    recipes: list[dict] = []
    ingredient_positions = [i for i, (line, _) in enumerate(parsed_lines) if re.match(r"^(Ingredientes|INGREDIENTES|ingredientes)\b", line, re.I)]

    for idx, ing_idx in enumerate(ingredient_positions):
        previous_modo = max(
            (j for j in range(ing_idx - 1, -1, -1) if re.match(r"^(Modo de Preparo|MODO DE PREPARO|Modo de preparo|Preparo|MODO)\b", parsed_lines[j][0], re.I)),
            default=-1,
        )

        if previous_modo >= 0:
            title_block = parsed_lines[previous_modo + 1 : ing_idx]
        else:
            current_category = parsed_lines[ing_idx][1]
            if current_category is not None:
                category_start = ing_idx
                while category_start > 0 and parsed_lines[category_start - 1][1] == current_category:
                    category_start -= 1
                title_block = parsed_lines[category_start:ing_idx]
            else:
                title_block = parsed_lines[:ing_idx]

        if not title_block:
            continue

        if previous_modo >= 0:
            split_point = max(
                (k for k, (line, _) in enumerate(title_block) if line.strip().endswith((".", "!", "?"))),
                default=-1,
            )
            title_block = title_block[split_point + 1 :]
            if not title_block:
                continue

        title_line = title_block[0][0]
        category = title_block[0][1]
        description = " ".join(line for line, _ in title_block[1:]).strip()

        modo_index = next(
            (j for j in range(ing_idx + 1, len(parsed_lines)) if re.match(r"^(Modo de Preparo|MODO DE PREPARO|Modo de preparo|Preparo|MODO)\b", parsed_lines[j][0], re.I)),
            None,
        )
        if modo_index is None:
            continue

        ingredients_lines = [parsed_lines[j][0] for j in range(ing_idx + 1, modo_index) if parsed_lines[j][0]]

        next_ing_idx = ingredient_positions[idx + 1] if idx + 1 < len(ingredient_positions) else len(parsed_lines)
        if next_ing_idx < len(parsed_lines):
            if next_ing_idx - 1 > modo_index and not re.match(r"^(Modo de Preparo|MODO DE PREPARO|Modo de preparo|Preparo|MODO)\b", parsed_lines[next_ing_idx - 1][0], re.I):
                next_recipe_title_index = next_ing_idx - 2 if next_ing_idx - 2 > modo_index else next_ing_idx - 1
            else:
                next_recipe_title_index = next_ing_idx - 1
            preparation_end = next_recipe_title_index - 1
        else:
            preparation_end = len(parsed_lines) - 1

        preparation_lines = [parsed_lines[j][0] for j in range(modo_index + 1, preparation_end + 1) if parsed_lines[j][0]]

        recipes.append(
            {
                "category": category or "Receitas",
                "recipe_name": title_line,
                "description": description,
                "ingredients_text": "\n".join(ingredients_lines).strip(),
                "preparation_text": "\n".join(preparation_lines).strip(),
            }
        )

    return recipes


def rewrite_introduction(author: str, publisher: str) -> str:
    return (
        f"Bem-vindo a '200 Receitas Low Carb', um guia completo de cozinha saudável elaborado por {author} e publicado pela {publisher}. "
        "Este livro oferece opções nutritivas para café da manhã, almoço, jantar, lanches e sobremesas, com foco em ingredientes de qualidade e equilíbrio. "
        "Os menus aqui apresentados foram pensados para facilitar sua rotina sem abrir mão do sabor, da variedade e da praticidade. "
        "Use este livro como ponto de partida para uma alimentação mais leve, inteligente e prazerosa, honrando seu compromisso com a saúde e o bem-estar."
    )


def insert_toc(document: Document) -> None:
    toc_paragraph = document.add_paragraph()
    fld_simple = OxmlElement("w:fldSimple")
    fld_simple.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    toc_paragraph._p.append(fld_simple)


def download_image(query: str, api_key: str, max_width: int = 1200) -> bytes | None:
    if not api_key:
        return None
    headers = {"Authorization": api_key}
    params = {"query": query, "per_page": 5, "orientation": "landscape", "size": "large"}
    response = requests.get("https://api.pexels.com/v1/search", headers=headers, params=params, timeout=20)
    if response.status_code != 200:
        return None
    data = response.json()
    photos = data.get("photos", [])
    if not photos:
        return None
    photo = photos[0]
    url = photo.get("src", {}).get("large") or photo.get("src", {}).get("original")
    if not url:
        return None
    image_resp = requests.get(url, timeout=30)
    if image_resp.status_code != 200:
        return None
    image = Image.open(BytesIO(image_resp.content))
    image = image.convert("RGB")
    width, height = image.size
    if width > max_width:
        ratio = max_width / float(width)
        new_height = int(height * ratio)
        image = image.resize((max_width, new_height), Image.LANCZOS)
    buffer = BytesIO()
    image.save(buffer, format="JPEG", quality=80, optimize=True)
    return buffer.getvalue()


def add_front_matter(document: Document, title: str, subtitle: str, author: str, publisher: str) -> None:
    document.add_paragraph(title, style="Title")
    document.add_paragraph(subtitle, style="Subtitle")
    document.add_paragraph(f"Autor: {author}", style="Intense Quote")
    document.add_paragraph(f"Editora: {publisher}", style="Intense Quote")
    document.add_page_break()

    document.add_paragraph("Direitos Autorais", style="Heading 1")
    document.add_paragraph(
        "© 2026 Mais Saúde. Todos os direitos reservados. Nenhuma parte desta obra pode ser reproduzida, armazenada em sistema de recuperação ou transmitida de qualquer forma ou por qualquer meio sem permissão por escrito do detentor dos direitos.")
    document.add_page_break()

    document.add_paragraph("Aviso Nutricional", style="Heading 1")
    document.add_paragraph(
        "As informações contidas neste livro são oferecidas apenas para fins educacionais e informativos. Cada pessoa pode reagir de forma diferente a ingredientes e receitas, e recomenda-se consultar um profissional de saúde ou nutricionista antes de fazer mudanças significativas na alimentação.")
    document.add_page_break()

    document.add_paragraph("Introdução", style="Heading 1")
    document.add_paragraph(rewrite_introduction(author, publisher))
    document.add_page_break()


def add_recipe_section(document: Document, recipe: dict, api_key: str, images_not_found: list[str]) -> None:
    category = recipe.get("category", "Receitas")
    recipe_name = recipe.get("recipe_name", "Receita")
    description = recipe.get("description", "")
    ingredients = recipe.get("ingredients_text", "")
    preparation = recipe.get("preparation_text", "")

    document.add_paragraph(category, style="Heading 1")
    document.add_paragraph(recipe_name, style="Heading 2")
    if description:
        document.add_paragraph(description)
    document.add_paragraph("Ingredientes e Modo de Preparo", style="Heading 3")

    if api_key:
        image_content = download_image(recipe_name, api_key)
        if image_content:
            image_stream = BytesIO(image_content)
            run = document.add_paragraph().add_run()
            run.add_picture(image_stream, width=Inches(6.5))
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            images_not_found.append(recipe_name)
    else:
        images_not_found.append(recipe_name)

    if ingredients:
        for paragraph in ingredients.split("\n"):
            document.add_paragraph(paragraph)

    if preparation:
        document.add_paragraph("")
        document.add_paragraph(preparation)

    document.add_page_break()


def write_support_files(output_dir: Path, description_text: str, review_entries: list[str], images_not_found: list[str]) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    (output_dir / "descricao-amazon.txt").write_text(description_text, encoding="utf-8")
    (output_dir / "relatorio-revisao.txt").write_text("\n".join(review_entries), encoding="utf-8")
    (output_dir / "imagens_nao_encontradas.txt").write_text("\n".join(images_not_found), encoding="utf-8")


def build_description(title: str, subtitle: str, author: str, publisher: str, recipe_count: int) -> str:
    return (
        f"{title} - {subtitle}\n"
        f"Autor: {author}\n"
        f"Editora: {publisher}\n\n"
        "Descubra 200 receitas low carb para uma alimentação saudável e equilibrada. "
        "Este livro reúne opções para café da manhã, almoço, jantar, lanches e sobremesas, com foco em sabor, praticidade e ingredientes acessíveis. "
        f"Em {recipe_count} receitas, você encontrará inspiração para todo o dia e apoio à sua jornada de saúde."
    )


def build_document(pdf_path: Path, output_docx: Path, api_key: str) -> tuple[list[str], list[str], int]:
    raw_text = extract_pdf_text(pdf_path)
    recipes = split_into_recipes(raw_text)
    if not recipes:
        raise RuntimeError("Não foi possível extrair receitas do PDF. Verifique o layout do arquivo ou faça ajustes no parser.")

    document = Document()
    add_front_matter(document, "200 Receitas Low Carb", "Café da manhã, almoço, jantar, lanches e sobremesas para uma alimentação saudável", "D.L. Martins", "Mais Saúde")
    insert_toc(document)
    document.add_page_break()

    images_not_found = []
    for recipe in recipes:
        add_recipe_section(document, recipe, api_key, images_not_found)

    document.add_paragraph("Sobre o Autor", style="Heading 1")
    document.add_paragraph(
        "D.L. Martins é autor especializado em culinária funcional e receitas saudáveis, com foco em rotinas acessíveis para famílias e quem busca uma vida mais equilibrada.")
    document.add_page_break()

    document.add_paragraph("Considerações Finais", style="Heading 1")
    document.add_paragraph(
        "Agradecemos por escolher este livro. Esperamos que cada receita inspire refeições cheias de sabor e bem-estar. Continue explorando sua cozinha low carb com curiosidade e carinho.")

    document.save(output_docx)

    review_entries = [
        f"Arquivo de origem: {pdf_path.name}",
        f"Arquivo gerado: {output_docx.name}",
        f"Receitas processadas: {len(recipes)}",
        f"Imagens não encontradas: {len(images_not_found)}",
        "Verifique no Word se o índice está atualizado e se todos os estilos de título foram aplicados corretamente.",
    ]
    return review_entries, images_not_found, len(recipes)


def main() -> None:
    parser = argparse.ArgumentParser(description="Gerador de livro KDP a partir de PDF de receitas low carb.")
    parser.add_argument("--pdf", required=True, help="Caminho para o PDF de origem.")
    parser.add_argument("--output", default="200-receitas-low-carb-dl-martins-kdp.docx", help="Nome do arquivo DOCX de saída.")
    parser.add_argument("--api-key", default=os.getenv("PEXELS_API_KEY"), help="Chave de API Pexels. Se não for informada, as imagens não serão procuradas.")
    args = parser.parse_args()

    pdf_path = Path(args.pdf)
    if not pdf_path.exists():
        print(f"Erro: arquivo PDF não encontrado em {pdf_path}", file=sys.stderr)
        sys.exit(1)

    output_docx_path = Path(args.output)
    output_dir = output_docx_path.parent or Path(".")
    review_entries, images_not_found, recipe_count = build_document(pdf_path, output_docx_path, args.api_key)

    description_text = build_description(
        "200 Receitas Low Carb",
        "Café da manhã, almoço, jantar, lanches e sobremesas para uma alimentação saudável",
        "D.L. Martins",
        "Mais Saúde",
        recipe_count,
    )
    write_support_files(output_dir, description_text, review_entries, images_not_found)

    print(f"Geração concluída: {output_docx_path}")
    print(f"Suporte criado em: {output_dir / 'descricao-amazon.txt'}, {output_dir / 'relatorio-revisao.txt'}, {output_dir / 'imagens_nao_encontradas.txt'}")


if __name__ == "__main__":
    main()
